import os
import subprocess
import sys

from PyQt5.QtWidgets import QApplication, QFileDialog

from utils.config_manager import get_printer_name, get_print_save_dir_for_today
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PrintCancelled(Exception):
    """用户取消打印时抛出"""
    pass


def print_doc(now_handle_data_id=-1, existing_file_path=None):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    def set_table_border(table, top=False, bottom=False, left=False, right=False, thickness=10):
        """设置表格边框粗细
        """
        # 获取表格属性
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        
        # 创建或获取表格边框设置
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is None:
            tbl_borders = OxmlElement('w:tblBorders')
            tbl_pr.append(tbl_borders)
        
        # 设置上下边框
        if top:
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), str(thickness))
            top_border.set(qn('w:space'), '0')
            top_border.set(qn('w:color'), 'auto')
            # 移除现有的top边框（如果存在）
            for e in tbl_borders.findall(qn('w:top')):
                tbl_borders.remove(e)
            tbl_borders.append(top_border)
            
        if bottom:
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), str(thickness))
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), 'auto')
            # 移除现有的bottom边框（如果存在）
            for e in tbl_borders.findall(qn('w:bottom')):
                tbl_borders.remove(e)
            tbl_borders.append(bottom_border)
        
        # 设置左右边框
        if left:
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), str(thickness))
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), 'auto')
            # 移除现有的left边框（如果存在）
            for e in tbl_borders.findall(qn('w:left')):
                tbl_borders.remove(e)
            tbl_borders.append(left_border)
        
        if right:
            right_border = OxmlElement('w:right')
            right_border.set(qn('w:val'), 'single')
            right_border.set(qn('w:sz'), str(thickness))
            right_border.set(qn('w:space'), '0')
            right_border.set(qn('w:color'), 'auto')
            # 移除现有的right边框（如果存在）
            for e in tbl_borders.findall(qn('w:right')):
                tbl_borders.remove(e)
            tbl_borders.append(right_border)
    
    import win32com.client

    def print_word_file(path, printer_name):
        word = win32com.client.Dispatch("Word.Application")
        word.ActivePrinter = printer_name
        doc = word.Documents.Open(path)
        # wdPrintFromTo=3，只打印第1页
        doc.PrintOut(Range=3, From="1", To="1")
        doc.Close(False)
        word.Quit()
    # 查询数据库获取数据
    if existing_file_path is not None:
        print_word_file(existing_file_path, get_printer_name())
        return
    
    from utils.data_manager import DataManager
    detail = DataManager.queryById(now_handle_data_id)
    # print(now_handle_data_id, detail)
    # 双重防护：无效记录或测试数据为空时直接返回，避免崩溃
    if detail is None:
        return
    x_list, y_list, highlight = DataManager.queryTestDataByFormId(now_handle_data_id)
    if not x_list or not y_list:
        return
    
    # 保存 Word 文件（使用配置的根目录/年份/月份/）
    save_dir = get_print_save_dir_for_today()
    # 出厂编号+测试时间，避免同出厂编号多次测试时文件名冲突
    # test_detail schema (data_manager.py):
    # 0 test_id
    # 1 project_name
    # 2 factory_number
    # 3 test_date
    # 4 pipeline_name
    # 5 line_hanger_no
    # 6 spec_model
    # 7 travel_direction
    # 8 working_load_n
    # ...
    test_time_safe = (str(detail[3]) if detail[3] else "").replace(":", "-").replace(" ", "_").replace("/", "-")
    time_suffix = f"-{test_time_safe}" if test_time_safe else ""
    filename = f"变力弹簧支吊架性能试验记录-{detail[2]}{time_suffix}.docx"
    full_path = os.path.abspath(os.path.join(save_dir, filename))

    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Title'].font.name = 'SimSun'
    doc.styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Heading 1'].font.name = 'SimSun'
    doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 获取第一个节（默认文档只有一个节）
    section = doc.sections[0]

    # 设置页面为 A4 纸（210mm × 297mm）
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # 设置页边距（单位：英寸），紧凑布局确保一页内
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.footer_distance = Inches(0.05)  # 页脚距底边，减小以压缩页脚区、缩短与正文的间距

    footer = section.footer
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run('第 1 页')
    run.font.size = Pt(6)


    # 添加公司名和标题（居中）
    def set_simsun_font(run, size=12, bold=False):
        """为指定的run设置宋体字体"""
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(size)
        run.bold = bold
        return run

    # 添加公司名和标题（居中）并设置宋体
    def add_centered_text_with_simsun(content, font_size=12, bold=False, style=None):
        paragraph = doc.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(content)
        set_simsun_font(run, font_size, bold)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
        # 设置段前段后为0，行距为单倍
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        return paragraph

    # 添加公司名（logo + 文字，同一行居中）
    title_para = doc.add_paragraph(style='Heading 1')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_para.paragraph_format.line_spacing = 1.0
    _res_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'resources')
    logo_path = os.path.join(_res_dir, 'logo.JPG') if os.path.exists(os.path.join(_res_dir, 'logo.JPG')) else os.path.join(_res_dir, 'logo.jpg')
    if os.path.exists(logo_path):
        run_logo = title_para.add_run()
        run_logo.add_picture(logo_path, width=Inches(0.8))
        title_para.add_run('  ')  # logo与文字间距
    run_text = title_para.add_run("江苏慧通管道设备股份有限公司")
    set_simsun_font(run_text, 24, bold=True)
    run_text.font.color.rgb = RGBColor(0, 0, 0)

    # 添加英文公司名（正常宋体）
    add_centered_text_with_simsun("Jiangsu Huitong PiPeline Equipment Co.Ltd.", font_size=14)

    # 添加主标题（大号宋体，加粗）
    add_centered_text_with_simsun("变力弹簧支吊架性能试验记录", font_size=20, bold=True, style='Heading 1')

    # 添加英文副标题（正常宋体）
    add_centered_text_with_simsun("Variable Force Spring Hanger Performance Test Record", font_size=12)

    # 添加基本信息表格
    table1 = doc.add_table(rows=3, cols=8)
    table1.style = 'Table Grid'
    col_widths = [Inches(1), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1),Inches(1)]
    for i, width in enumerate(col_widths):
        for row in table1.rows:
            row.cells[i].width = width

    _r0 = table1.rows[0]
    _r0.cells[1].merge(_r0.cells[2])
    _r0.cells[1].merge(_r0.cells[2])

    # detail 字段顺序见 DataManager.TEST_DETAIL_SELECT_COLS（test_detail）
    # 第一行内容 工程名称 出厂编号 试验日期
    table1.rows[0].cells[0].text = "工程名称"
    table1.rows[0].cells[4].text = "出厂编号"
    table1.rows[0].cells[6].text = "试验日期"
    # 第二行内容 管系名称 管线号-支吊点号 规格型号 位移方向(+/-)
    table1.rows[1].cells[0].text = "管系名称"
    table1.rows[1].cells[2].text = "管线号-支吊点号"
    table1.rows[1].cells[4].text = "规格型号"
    table1.rows[1].cells[6].text = "位移方向(+/-)"
    # 第三行内容 工作/热态载荷 安装/冷态载荷 安装/冷态位置 螺纹尺寸
    table1.rows[2].cells[0].text = "工作/热态载荷(N)"
    table1.rows[2].cells[2].text = "安装/冷态载荷(N)"
    table1.rows[2].cells[4].text = "安装/冷态位置(mm)"
    table1.rows[2].cells[6].text = "螺纹尺寸(M)"

    def safe_str(val):
        return "" if val is None else str(val)

    # detail[0]=test_id, [1]工程名称…[21]测试结论, [22]=file_path（见 data_manager.TEST_DETAIL_COLUMNS）
    if detail:
        table1.rows[0].cells[1].text = safe_str(detail[1])
        table1.rows[0].cells[5].text = safe_str(detail[2])
        table1.rows[0].cells[7].text = safe_str(detail[3])

        table1.rows[1].cells[1].text = safe_str(detail[4])
        table1.rows[1].cells[3].text = safe_str(detail[5])
        table1.rows[1].cells[5].text = safe_str(detail[6])
        table1.rows[1].cells[7].text = safe_str(detail[7])

        table1.rows[2].cells[1].text = safe_str(detail[8])
        table1.rows[2].cells[3].text = safe_str(detail[9])
        table1.rows[2].cells[5].text = safe_str(detail[10])
        table1.rows[2].cells[7].text = safe_str(detail[11])
    else:
        table1.rows[0].cells[1].text = ""
        table1.rows[0].cells[5].text = ""
        table1.rows[0].cells[7].text = ""
        table1.rows[1].cells[1].text = ""
        table1.rows[1].cells[3].text = ""
        table1.rows[1].cells[5].text = ""
        table1.rows[1].cells[7].text = ""
        table1.rows[2].cells[1].text = ""
        table1.rows[2].cells[3].text = ""
        table1.rows[2].cells[5].text = ""
        table1.rows[2].cells[7].text = ""
    format_table_cells(table1, font_size=10)
    # 设置第一个表格的上边和左右边加粗
    set_table_border(table1, top=True, left=True, right=True)


    # 添加一个 1x1 的表格
    table2 = doc.add_table(rows=1, cols=1)
    table2.cell(0,0).width = Inches(8)
    table2.style = 'Table Grid'
    # 设置这个表格的左右边加粗
    set_table_border(table2, left=True, right=True)
    # 获取单元格
    cell = table2.cell(0, 0)
    # 在单元格内添加图片
    paragraph = cell.paragraphs[0]  # 获取单元格内的段落
    picture = paragraph.add_run().add_picture('./resources/png.png', width=Inches(7))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(4)
    paragraph_format.space_after = Pt(4)  # 减小间距，避免内容溢出产生空白第二页

    # 第三个表格为6*8
    table3 = doc.add_table(rows=6, cols=8)
    table3.style = 'Table Grid'
    col_widths = [Inches(1), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1),Inches(1)]
    for i, width in enumerate(col_widths):
        for row in table3.rows:
            row.cells[i].width = width
    
    table3.cell(0,0).merge(table3.cell(1,0))
    table3.cell(0,1).merge(table3.cell(1,1))
    table3.cell(0,2).merge(table3.cell(1,2))
    table3.cell(0,3).merge(table3.cell(1,3))
    table3.cell(0,4).merge(table3.cell(1,4))
    table3.cell(2,0).merge(table3.cell(3,0))
    table3.cell(2,1).merge(table3.cell(3,1))
    table3.cell(2,2).merge(table3.cell(3,2))
    table3.cell(2,3).merge(table3.cell(3,3))
    table3.cell(2,4).merge(table3.cell(3,4))
    _tr4 = table3.rows[4]
    _tr4.cells[3].merge(_tr4.cells[4])
    _tr4.cells[3].merge(_tr4.cells[4])
    _tr5 = table3.rows[5]
    _tr5.cells[1].merge(_tr5.cells[2])
    _tr5.cells[1].merge(_tr5.cells[2])

    # 第一行包括：整定载荷实测值(N)、载荷偏差度、弹簧刚度、理论值、实测值、刚度偏差
    table3.cell(0,0).text = "整定载荷实测值(N)"
    table3.cell(0,2).text = "载荷偏差度"
    table3.cell(0,4).text = "弹簧刚度"
    table3.cell(0,5).text = "理论值"
    table3.cell(0,6).text = "实测值"
    table3.cell(0,7).text = "刚度偏差"
    # 第二行包括：最小位移(mm)、最小位移相应的实测载荷(N)、超载试验、超载试验值、超载起始-终止时间、超载保持时间
    table3.cell(2,0).text = "最小位移(mm)"
    table3.cell(2,2).text = "最小位移相应的实测载荷(N)"
    table3.cell(2,4).text = "超载试验"
    table3.cell(2,5).text = "超载试验值"
    table3.cell(2,6).text = "超载起始-终止时间"
    table3.cell(2,7).text = "超载保持时间"
    # 第三行包括：最大位移(mm)、最大位移相应的实测载荷(N)、测试结论
    table3.cell(4,0).text = "最大位移(mm)"
    table3.cell(4,2).text = "最大位移相应的实测载荷(N)"
    table3.cell(4,6).text = "测试结论"
    # 第四行包括：试验人员、批准人员
    table3.cell(5,0).text = "试验人员"
    table3.cell(5,5).text = "批准人员"

    # 弹簧刚度：库中仅 spring_stiffness 一列，理论值/实测值均填该值；刚度偏差暂无独立字段
    # 超载相关：test_detail 无对应列，留空
    if detail:
        table3.rows[0].cells[1].text = safe_str(detail[15])
        table3.rows[0].cells[3].text = safe_str(detail[16])
        table3.rows[1].cells[5].text = safe_str(detail[12])
        table3.rows[1].cells[6].text = safe_str(detail[12])
        table3.rows[1].cells[7].text = ""
        table3.rows[2].cells[1].text = safe_str(detail[17])
        table3.rows[2].cells[3].text = safe_str(detail[18])
        table3.rows[3].cells[5].text = ""
        table3.rows[3].cells[6].text = ""
        table3.rows[3].cells[7].text = ""
        table3.rows[4].cells[1].text = safe_str(detail[19])
        table3.rows[4].cells[3].text = safe_str(detail[20])
        table3.rows[4].cells[7].text = safe_str(detail[21])
        table3.rows[5].cells[1].text = safe_str(detail[13])
        table3.rows[5].cells[6].text = safe_str(detail[14])
    else:
        table3.rows[0].cells[1].text = ""
        table3.rows[0].cells[3].text = ""
        table3.rows[1].cells[5].text = ""
        table3.rows[1].cells[6].text = ""
        table3.rows[1].cells[7].text = ""
        table3.rows[2].cells[1].text = ""
        table3.rows[2].cells[3].text = ""
        table3.rows[3].cells[5].text = ""
        table3.rows[3].cells[6].text = ""
        table3.rows[3].cells[7].text = ""
        table3.rows[4].cells[1].text = ""
        table3.rows[4].cells[3].text = ""
        table3.rows[4].cells[7].text = ""
        table3.rows[5].cells[1].text = ""
        table3.rows[5].cells[6].text = ""

    format_table_cells(table3, font_size=10)
    # 设置这个表格的左右边加粗
    set_table_border(table3, left=True, right=True)

    doc.save(full_path)

    # 将文件完整路径写入数据库
    DataManager.update_file_path(now_handle_data_id, full_path)
    print_word_file(full_path, get_printer_name())


def format_table_cells(table, font_size=9):
    for row in table.rows:
        for cell in row.cells:
            # 设置垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for paragraph in cell.paragraphs:
                # 设置水平居中
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


import pypandoc

def convert_docx_to_pdf(input_path, output_path):
    output = pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
    return output_path if output == "" else None


if __name__ == '__main__':
    # pdf_path = convert_docx_to_pdf("恒力吊架性能试验记录（无图）.docx", "output.pdf")
    # print("转换成功：" + pdf_path)
    print_doc()
